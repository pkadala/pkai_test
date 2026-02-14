#!/usr/bin/env python3
"""Build docs/workspace-mcp-env.docx from the markdown content (readable Word format)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str) -> None:
    """Set table cell background color (e.g. 'E8F4EA' for light green)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D9E2F3") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hrow = table.rows[0].cells
    for i, h in enumerate(headers):
        hrow[i].text = h
        set_cell_shading(hrow[i], header_fill)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            if c < len(table.rows[r + 1].cells):
                table.rows[r + 1].cells[c].text = val
    doc.add_paragraph()


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    out_path = root / "docs" / "workspace-mcp-env.docx"
    root.mkdir(exist_ok=True)
    (root / "docs").mkdir(exist_ok=True)

    doc = Document()
    doc.add_heading("Workspace MCP Configuration", 0)

    p = doc.add_paragraph()
    p.add_run(
        "This document lists the environment variables used for the workspace-mcp integration "
        "(Google Drive & Tasks) in stdio vs streamable-http mode, and which apply to the app server vs the MCP server."
    )

    # --- Stdio mode ---
    add_heading(doc, "Stdio mode (single app server)", level=1)
    add_paragraph(
        doc,
        "The app spawns workspace-mcp as a subprocess. There is no separate MCP server. All variables are set on the app server."
    )
    add_table(
        doc,
        ["Variable", "Required?", "Purpose"],
        [
            ["WORKSPACE_MCP_TRANSPORT", "Optional", "Omit or set to stdio (default)."],
            [
                "WORKSPACE_MCP_OAUTH_REDIRECT_URI",
                "Yes",
                "Full OAuth redirect URI for workspace-mcp (e.g. http://localhost:8765/oauth2callback). The app derives port and base URI from this and passes them to the subprocess.",
            ],
            ["GOOGLE_OAUTH_CLIENT_ID", "Yes (for OAuth)", "Passed to the subprocess so workspace-mcp can complete Google OAuth."],
            ["GOOGLE_OAUTH_CLIENT_SECRET", "Yes (for OAuth)", "Same."],
            ["USER_GOOGLE_EMAIL", "Optional", "Passed to the subprocess; used by some tools."],
            ["WORKSPACE_MCP_COMMAND", "Optional", "Override executable (default: workspace-mcp or python -m workspace_mcp)."],
            ["WORKSPACE_MCP_ARGS", "Optional", "Override arguments (default: --tools drive tasks)."],
        ],
    )
    add_paragraph(doc, "Not needed on app server for stdio: WORKSPACE_MCP_HTTP_URL.", bold=True)

    # --- Streamable-HTTP: App server ---
    add_heading(doc, "Streamable-HTTP mode (app server + MCP server)", level=1)
    add_paragraph(
        doc,
        "The app connects to a separate workspace-mcp service over HTTP. Configure two places: the app server and the MCP server."
    )
    add_heading(doc, "App server (PKAI / FastAPI)", level=2)
    add_table(
        doc,
        ["Variable", "Required?", "Purpose"],
        [
            ["WORKSPACE_MCP_TRANSPORT", "Yes", "Must be streamable-http."],
            ["WORKSPACE_MCP_HTTP_URL", "Yes", "Full MCP endpoint URL (e.g. https://mcp-service.up.railway.app/mcp)."],
        ],
    )
    add_paragraph(
        doc,
        "Not needed on app server for streamable-http: WORKSPACE_MCP_OAUTH_REDIRECT_URI, GOOGLE_OAUTH_CLIENT_ID, "
        "GOOGLE_OAUTH_CLIENT_SECRET, USER_GOOGLE_EMAIL — the app does not use or pass these when using streamable-http.",
        bold=True,
    )

    # --- Streamable-HTTP: MCP server ---
    add_heading(doc, "MCP server (separate workspace-mcp process)", level=2)
    add_paragraph(
        doc,
        "Set these in the environment of the process that runs workspace-mcp --transport streamable-http (e.g. a second Railway service)."
    )
    add_table(
        doc,
        ["Variable", "Required?", "Purpose"],
        [
            ["PORT or WORKSPACE_MCP_PORT", "Yes", "Port the server listens on (Railway often sets PORT)."],
            ["WORKSPACE_MCP_BASE_URI", "Yes", "Public base URL of this MCP service (e.g. https://mcp-service.up.railway.app)."],
            ["GOOGLE_OAUTH_REDIRECT_URI", "Yes", "Full redirect URI (e.g. https://mcp-service.up.railway.app/oauth2callback). Must match Google Cloud Console."],
            ["GOOGLE_OAUTH_CLIENT_ID", "Yes", "Google OAuth client ID."],
            ["GOOGLE_OAUTH_CLIENT_SECRET", "Yes", "Google OAuth client secret."],
            ["USER_GOOGLE_EMAIL", "Optional", "For tools that need a user email."],
        ],
    )

    # --- Quick reference ---
    add_heading(doc, "Quick reference", level=1)
    add_table(
        doc,
        ["Variable", "Stdio (app server)", "Streamable-http (app server)", "Streamable-http (MCP server)"],
        [
            ["WORKSPACE_MCP_TRANSPORT", "Optional (default stdio)", "Required (streamable-http)", "—"],
            ["WORKSPACE_MCP_HTTP_URL", "Not used", "Required", "—"],
            ["WORKSPACE_MCP_OAUTH_REDIRECT_URI", "Required", "Not used", "Required (or equivalent)"],
            ["GOOGLE_OAUTH_CLIENT_ID", "Required (for OAuth)", "Not used", "Required"],
            ["GOOGLE_OAUTH_CLIENT_SECRET", "Required (for OAuth)", "Not used", "Required"],
            ["USER_GOOGLE_EMAIL", "Optional", "Not used", "Optional"],
            ["PORT / WORKSPACE_MCP_PORT", "—", "—", "Required"],
            ["WORKSPACE_MCP_BASE_URI", "Derived from redirect URI", "Not used", "Required"],
            ["WORKSPACE_MCP_COMMAND / WORKSPACE_MCP_ARGS", "Optional", "Not used", "Optional"],
        ],
    )

    # --- Local vs production ---
    add_heading(doc, "Local vs production (Railway)", level=1)
    bullets = [
        "Local (stdio): Use WORKSPACE_MCP_OAUTH_REDIRECT_URI=http://localhost:8765/oauth2callback. workspace-mcp runs on the same machine and can open a server on port 8765 for the OAuth callback.",
        "Railway (stdio): A single app process cannot reliably run workspace-mcp's OAuth server (port conflict). Prefer streamable-http with a separate MCP server that has its own public URL and handles the OAuth callback.",
        "Railway (streamable-http): Deploy workspace-mcp as a second service; set its public URL as WORKSPACE_MCP_BASE_URI and GOOGLE_OAUTH_REDIRECT_URI, and add that redirect URI in Google Cloud Console. On the app service, set only WORKSPACE_MCP_TRANSPORT=streamable-http and WORKSPACE_MCP_HTTP_URL.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
